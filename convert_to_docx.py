#!/usr/bin/env python3
"""Convert markdown to Word document for FDHC 1099 Sales Pay Plans"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_custom(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 153)
        else:
            run.font.size = Pt(12)
            run.font.bold = True
    return heading

def add_paragraph_custom(doc, text, bold=False, italic=False, indent=False):
    """Add a paragraph with custom formatting"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_table_custom(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
    
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
    
    return table

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Factory Direct Homes Center', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading('1099 Independent Contractor Sales Compensation Plans', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 102, 153)
    
    location = doc.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = location.add_run('Auburn, Indiana | DeKalb County')
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.italic = True
    
    doc.add_paragraph()
    
    # Effective Date
    add_paragraph_custom(doc, 'Effective Date: [Insert Date]', bold=True)
    add_paragraph_custom(doc, 'Classification: Independent Contractor (Form 1099-NEC)', bold=True)
    doc.add_paragraph()
    
    # Critical Legal Framework
    add_heading_custom(doc, 'CRITICAL LEGAL FRAMEWORK & COMPLIANCE', level=1)
    
    add_heading_custom(doc, 'IRS Classification Standards (3-Part Test)', level=2)
    
    add_heading_custom(doc, '1. Behavioral Control', level=3)
    add_paragraph_custom(doc, 'PROHIBITED: Directing WHEN, WHERE, or HOW work is performed')
    add_paragraph_custom(doc, 'PROHIBITED: Requiring specific hours, attendance, or location')
    add_paragraph_custom(doc, 'PROHIBITED: Mandating use of specific scripts, methods, or processes')
    add_paragraph_custom(doc, 'ALLOWED: Specifying RESULTS desired (sales closed, revenue generated)')
    add_paragraph_custom(doc, 'ALLOWED: Setting performance standards based on outcomes')
    add_paragraph_custom(doc, 'ALLOWED: Providing product information and pricing guidelines')
    
    add_heading_custom(doc, '2. Financial Control', level=3)
    add_paragraph_custom(doc, 'REQUIRED: Contractor must have opportunity for profit OR loss')
    add_paragraph_custom(doc, 'REQUIRED: Contractor pays their own business expenses')
    add_paragraph_custom(doc, 'REQUIRED: No guaranteed minimum pay, salary, or draws')
    add_paragraph_custom(doc, 'ALLOWED: Commission-only compensation')
    add_paragraph_custom(doc, 'ALLOWED: Performance bonuses tied to results')
    add_paragraph_custom(doc, 'PROHIBITED: Reimbursing routine business expenses (creates employee-like relationship)')
    
    add_heading_custom(doc, '3. Type of Relationship', level=3)
    add_paragraph_custom(doc, 'REQUIRED: Written Independent Contractor Agreement')
    add_paragraph_custom(doc, 'REQUIRED: No employee benefits (health insurance, 401k, PTO)')
    add_paragraph_custom(doc, 'ALLOWED: Project-based or indefinite term (but indefinite suggests employee status)')
    add_paragraph_custom(doc, 'HIGH RISK: Sales work is integral to FDHC\'s core business (this weighs toward employee classification)')
    
    doc.add_page_break()
    
    # DOL Economic Reality Test
    add_heading_custom(doc, 'DOL Economic Reality Test (2024 Final Rule - 29 CFR Part 795)', level=2)
    add_paragraph_custom(doc, 'Six Factors (no single factor is determinative):', bold=True)
    
    add_paragraph_custom(doc, '1. Opportunity for profit or loss depending on managerial skill - Contractor must be able to affect earnings through their own business decisions. Must be able to hire helpers, advertise, choose which leads to pursue.')
    add_paragraph_custom(doc, '2. Investments by the worker and the employer - Contractor should invest in their own business (marketing, tools, vehicle). FDHC investment in contractor suggests employee relationship.')
    add_paragraph_custom(doc, '3. Permanence of the work relationship - Indefinite, continuous relationship suggests employment. Project-based or definite term suggests contractor.')
    add_paragraph_custom(doc, '4. Nature and degree of control - Less control = more likely contractor. Control over scheduling, methods, territory suggests employee.')
    add_paragraph_custom(doc, '5. Whether the work performed is integral to the employer\'s business - Sales IS integral to FDHC (HIGH RISK FACTOR). This factor alone does not determine status.')
    add_paragraph_custom(doc, '6. Skill and initiative - Specialized skills and independent business judgment suggest contractor. Use of initiative to secure business opportunities.')
    
    add_heading_custom(doc, 'Indiana-Specific Considerations', level=2)
    add_paragraph_custom(doc, 'Indiana Department of Labor (IDOL):')
    add_paragraph_custom(doc, '• Indiana follows federal IRS guidelines for classification', indent=True)
    add_paragraph_custom(doc, '• No state-specific "ABC test" like California or New Jersey', indent=True)
    add_paragraph_custom(doc, '• Worker misclassification penalties apply under Indiana Code 22-1-1-1 et seq.', indent=True)
    
    add_paragraph_custom(doc, 'Indiana Non-Compete Law (IC 25-5-3):')
    add_paragraph_custom(doc, '• Non-compete agreements are generally enforceable in Indiana', indent=True)
    add_paragraph_custom(doc, '• Must be reasonable in scope, duration, and geographic area', indent=True)
    add_paragraph_custom(doc, '• Independent contractors CAN be subject to non-competes', indent=True)
    add_paragraph_custom(doc, '• Courts scrutinize more heavily for low-wage workers', indent=True)
    
    add_paragraph_custom(doc, 'Indiana Tax Requirements:')
    add_paragraph_custom(doc, '• Form 1099-NEC required for payments ≥$600/year', indent=True)
    add_paragraph_custom(doc, '• No state income tax withholding required for 1099 contractors', indent=True)
    add_paragraph_custom(doc, '• Contractor responsible for self-employment tax (15.3%)', indent=True)
    
    doc.add_page_break()
    
    # High-Risk Provisions
    add_heading_custom(doc, 'HIGH-RISK PROVISIONS TO AVOID', level=1)
    add_paragraph_custom(doc, 'DO NOT INCLUDE:', bold=True)
    
    risks = [
        '1. Guaranteed Minimum Pay - Any draw, advance, or guaranteed amount triggers W-2',
        '2. Required Hours - Cannot mandate 9-5, specific days, or attendance',
        '3. Exclusive Employment - Cannot prohibit working for competitors',
        '4. Expense Reimbursement - Paying for gas, phone, meals suggests employee',
        '5. Benefits - Health insurance, retirement, PTO are employee-only',
        '6. Detailed Scripts - Can provide info, cannot mandate specific sales methods',
        '7. Territory Restrictions - Telling contractor WHERE to work suggests control',
        '8. Sales Quotas with Penalties - Can set goals, cannot punish for missing them',
        '9. Required Training on Methods - Product training OK, sales technique training = risk',
        '10. Company Email/Phone Required - Providing tools suggests employee'
    ]
    for risk in risks:
        add_paragraph_custom(doc, risk)
    
    add_paragraph_custom(doc, 'SAFE PRACTICES:', bold=True)
    
    safe = [
        '1. Commission-Only Pay - No base, no guarantee, no draw',
        '2. Results-Based Metrics - Pay only when sales close',
        '3. Flexible Scheduling - Contractor sets own hours',
        '4. Multiple Clients Allowed - Explicitly permit other sales work',
        '5. Contractor-Paid Expenses - They bear cost of doing business',
        '6. Written Agreement - Clear IC classification with all parties signing',
        '7. 1099-NEC Filing - Report all payments annually',
        '8. E&O Insurance Requirement - Shift liability to contractor',
        '9. Hold Harmless Clause - Contractual liability protection',
        '10. Indemnification - Contractor indemnifies FDHC for their actions'
    ]
    for s in safe:
        add_paragraph_custom(doc, s)
    
    doc.add_page_break()
    
    # Plan 1
    add_heading_custom(doc, 'PLAN VARIATION 1: STRAIGHT COMMISSION MODEL', level=1)
    add_paragraph_custom(doc, 'Overview:', bold=True)
    add_paragraph_custom(doc, 'Pure commission structure with no base pay, no guarantees, and payment only upon successful closing and funding.')
    
    add_paragraph_custom(doc, 'Commission Structure:', bold=True)
    add_table_custom(doc, 
        ['Home Type', 'Commission Rate', 'Calculation Base'],
        [
            ['Single Wide', '3%', 'Net Sale Price (after discounts)'],
            ['Double Wide', '3%', 'Net Sale Price (after discounts)'],
            ['Modular', '2.5%', 'Net Sale Price (after discounts)'],
            ['Land-Home Package', '2%', 'Total Package Price']
        ]
    )
    
    doc.add_paragraph()
    add_paragraph_custom(doc, 'Down Payment Compensation:', bold=True)
    
    add_paragraph_custom(doc, 'Option A: Split Commission on Down Payment', bold=True)
    add_paragraph_custom(doc, '• Contractor receives 50% of down payment amount as advance against total commission')
    add_paragraph_custom(doc, '• Example: $10,000 down payment = $5,000 paid to contractor within 5 business days')
    add_paragraph_custom(doc, '• Remaining commission paid at closing/funding')
    add_paragraph_custom(doc, '• RISK LEVEL: MEDIUM - IRS may view as constructive receipt')
    
    add_paragraph_custom(doc, 'Option B: Referral Fee Structure (SAFER)', bold=True)
    add_paragraph_custom(doc, '• Contractor receives flat referral fee when customer makes down payment: $500-$1,000')
    add_paragraph_custom(doc, '• This is separate from sales commission')
    add_paragraph_custom(doc, '• Requires customer to sign referral disclosure')
    add_paragraph_custom(doc, '• RISK LEVEL: LOW - Clearly separate transaction')
    
    add_paragraph_custom(doc, 'Option C: Milestone Payment (RECOMMENDED)', bold=True)
    add_paragraph_custom(doc, '• 25% of total commission paid upon signed purchase agreement + down payment')
    add_paragraph_custom(doc, '• 75% paid at closing/funding')
    add_paragraph_custom(doc, '• Documented as "performance-based milestone" not "advance"')
    add_paragraph_custom(doc, '• RISK LEVEL: LOW - Tied to binding customer commitment')
    
    doc.add_paragraph()
    add_paragraph_custom(doc, 'Payment Timeline:', bold=True)
    add_table_custom(doc,
        ['Event', 'Payment', 'Timing'],
        [
            ['Signed PA + Down Payment', '25% of commission', 'Within 10 business days'],
            ['Home Delivered', '25% of commission', 'Within 10 business days'],
            ['Final Inspection Complete', '25% of commission', 'Within 10 business days'],
            ['Funding/Closing', '25% of commission', 'Within 5 business days']
        ]
    )
    
    doc.add_page_break()
    
    add_paragraph_custom(doc, 'Chargebacks & Clawbacks:', bold=True)
    add_paragraph_custom(doc, '• If sale cancels before delivery: 100% of paid commission clawed back')
    add_paragraph_custom(doc, '• If sale cancels after delivery: 50% of paid commission clawed back')
    add_paragraph_custom(doc, '• If financing falls through: 100% clawback of advance payments')
    add_paragraph_custom(doc, '• Contractor has 30 days to repay or amount deducted from future commissions')
    
    add_paragraph_custom(doc, 'Definitions:', bold=True)
    add_paragraph_custom(doc, '"Net Sale Price" - The final purchase price after all customer discounts, promotions, and credits, excluding taxes, title, delivery, and setup fees.')
    add_paragraph_custom(doc, '"Closing/Funding" - The date on which FDHC receives full payment from customer, lender, or third-party financier.')
    add_paragraph_custom(doc, '"Signed Purchase Agreement" - A fully executed purchase agreement with customer signature, contractor signature, and all required deposits.')
    add_paragraph_custom(doc, '"Down Payment" - Non-refundable customer deposit of at least $1,000 or 10% of purchase price, whichever is greater.')
    
    doc.add_page_break()
    
    # Plan 2
    add_heading_custom(doc, 'PLAN VARIATION 2: TIERED PERFORMANCE MODEL', level=1)
    add_paragraph_custom(doc, 'Overview:', bold=True)
    add_paragraph_custom(doc, 'Accelerated commission rates based on volume and margin, rewarding high performers while maintaining pure 1099 classification.')
    
    add_paragraph_custom(doc, 'Base Commission Rates:', bold=True)
    add_table_custom(doc,
        ['Monthly Volume', 'Commission Rate', 'Bonus'],
        [
            ['0-2 units', '2.5%', 'None'],
            ['3-4 units', '3.0%', '$250/unit'],
            ['5+ units', '3.5%', '$500/unit']
        ]
    )
    
    doc.add_paragraph()
    add_paragraph_custom(doc, 'Margin Multiplier:', bold=True)
    add_table_custom(doc,
        ['Gross Margin', 'Multiplier'],
        [
            ['Under 15%', '0.75x'],
            ['15-20%', '1.0x'],
            ['20-25%', '1.25x'],
            ['Over 25%', '1.5x']
        ]
    )
    
    doc.add_paragraph()
    add_paragraph_custom(doc, 'Example Calculation:', bold=True)
    add_paragraph_custom(doc, '• Sale Price: $100,000')
    add_paragraph_custom(doc, '• Volume Tier: 3 units (3.0% base)')
    add_paragraph_custom(doc, '• Gross Margin: 22% (1.25x multiplier)')
    add_paragraph_custom(doc, '• Commission: $100,000 × 0.03 × 1.25 = $3,750 + $250 bonus = $4,000')
    
    doc.add_paragraph()
    add_paragraph_custom(doc, 'Down Payment Structure:', bold=True)
    add_paragraph_custom(doc, '"Customer Commitment Bonus" (1099-Safe Terminology)', bold=True)
    
    add_table_custom(doc,
        ['Down Payment Amount', 'Contractor Payment'],
        [
            ['$1,000 - $2,499', '$250'],
            ['$2,500 - $4,999', '$500'],
            ['$5,000 - $9,999', '$750'],
            ['$10,000+', '$1,000']
        ]
    )
    
    doc.add_paragraph()
    add_paragraph_custom(doc, 'Terms:', bold=True)
    add_paragraph_custom(doc, '• Paid within 5 business days of down payment clearing')
    add_paragraph_custom(doc, '• Labeled as "Customer Acquisition Bonus" not "commission advance"')
    add_paragraph_custom(doc, '• Separate 1099-MISC box 3 (Other Income) or included in 1099-NEC')
    add_paragraph_custom(doc, '• Clawed back if sale cancels within 30 days')
    
    doc.add_page_break()
    
    add_paragraph_custom(doc, 'Self-Generated Lead Bonus:', bold=True)
    add_paragraph_custom(doc, '• Contractor receives additional 0.5% if lead is self-generated (not FDHC-provided)')
    add_paragraph_custom(doc, '• Must document lead source at time of first contact')
    add_paragraph_custom(doc, '• Cannot use FDHC CRM, website, or phone system to generate leads')
    add_paragraph_custom(doc, '• Contractor must maintain own lead generation systems')
    
    add_paragraph_custom(doc, 'Payment Schedule:', bold=True)
    add_table_custom(doc,
        ['Milestone', 'Payment %', 'Trigger'],
        [
            ['Contract Signed', '20%', 'PA executed, down payment received'],
            ['Credit Approved', '20%', 'Financing commitment obtained'],
            ['Home Ordered', '20%', 'Factory order placed'],
            ['Home Delivered', '20%', 'Physical delivery complete'],
            ['Final Funding', '20%', 'All funds received by FDHC']
        ]
    )
    
    doc.add_paragraph()
    add_paragraph_custom(doc, 'Annual Performance Bonus (1099-Safe):', bold=True)
    add_table_custom(doc,
        ['Annual Units Sold', 'Year-End Bonus'],
        [
            ['10-14 units', '$2,500'],
            ['15-19 units', '$5,000'],
            ['20-24 units', '$10,000'],
            ['25+ units', '$15,000 + 0.5% override on all sales']
        ]
    )
    
    add_paragraph_custom(doc, 'Key: This is a year-end performance bonus, NOT a guaranteed payment. Contractor must achieve threshold to qualify.')
    
    doc.add_page_break()
    
    # Plan 3
    add_heading_custom(doc, 'PLAN VARIATION 3: DRAW AGAINST COMMISSION MODEL', level=1)
    
    # Warning box
    warning_para = doc.add_paragraph()
    warning_run = warning_para.add_run('⚠️ LEGAL WARNING')
    warning_run.bold = True
    warning_run.font.size = Pt(14)
    warning_run.font.color.rgb = RGBColor(255, 0, 0)
    
    add_paragraph_custom(doc, 'This model carries elevated IRS reclassification risk. Only use with extensive legal documentation and clear understanding that "draws" are actually advances against future commissions, not guaranteed pay.')
    
    add_paragraph_custom(doc, 'Structure:', bold=True)
    add_paragraph_custom(doc, 'Monthly Advance ("Draw"):', bold=True)
    add_paragraph_custom(doc, '• $2,000/month advance against future commissions')
    add_paragraph_custom(doc, '• NOT a salary or guaranteed payment')
    add_paragraph_custom(doc, '• Must be repaid from future commissions')
    add_paragraph_custom(doc, '• If commissions don\'t cover draw, contractor owes balance')
    
    add_paragraph_custom(doc, 'Commission Rates:', bold=True)
    add_paragraph_custom(doc, '• Single/Double Wide: 4%')
    add_paragraph_custom(doc, '• Modular: 3%')
    add_paragraph_custom(doc, '• Minimum commission per sale: $2,000')
    
    add_paragraph_custom(doc, 'Draw Mechanics:', bold=True)
    add_paragraph_custom(doc, 'Month 1 Example:')
    add_paragraph_custom(doc, '• Draw taken: $2,000')
    add_paragraph_custom(doc, '• Commissions earned: $0')
    add_paragraph_custom(doc, '• Balance owed: $2,000', indent=True)
    
    add_paragraph_custom(doc, 'Month 2 Example:')
    add_paragraph_custom(doc, '• Draw taken: $2,000')
    add_paragraph_custom(doc, '• Commissions earned: $6,000')
    add_paragraph_custom(doc, '• Less: Prior month draw repayment: $2,000')
    add_paragraph_custom(doc, '• Less: Current month draw repayment: $2,000')
    add_paragraph_custom(doc, '• Net payment: $2,000', indent=True)
    
    doc.add_page_break()
    
    add_paragraph_custom(doc, 'Draw Agreement Requirements (CRITICAL):', bold=True)
    add_paragraph_custom(doc, 'To maintain 1099 status with a draw structure, the agreement MUST include:')
    add_paragraph_custom(doc, '1. Promissory Note Character - Draws are loans, not wages')
    add_paragraph_custom(doc, '2. Repayment Obligation - Contractor legally obligated to repay')
    add_paragraph_custom(doc, '3. No Minimum Hours - Draw available regardless of hours worked')
    add_paragraph_custom(doc, '4. No Performance Requirements - Draw not tied to activity metrics')
    add_paragraph_custom(doc, '5. Interest Charge - Nominal interest (e.g., 2% annually) to reinforce loan nature')
    add_paragraph_custom(doc, '6. Separate Accounting - Draws tracked separately from commissions')
    
    add_paragraph_custom(doc, 'Down Payment Treatment:', bold=True)
    add_paragraph_custom(doc, 'NO down payment advances with draw model - Double advance creates significant employee classification risk.')
    add_paragraph_custom(doc, 'Instead:')
    add_paragraph_custom(doc, '• Full commission paid at closing only')
    add_paragraph_custom(doc, '• OR milestone payments without draw structure')
    
    add_paragraph_custom(doc, 'Termination Provisions:', bold=True)
    add_paragraph_custom(doc, 'Upon termination (by either party):')
    add_paragraph_custom(doc, '• All outstanding draw balances become immediately due')
    add_paragraph_custom(doc, '• FDHC may withhold final commissions to satisfy draw balance')
    add_paragraph_custom(doc, '• Contractor remains liable for any deficiency')
    add_paragraph_custom(doc, '• Collection costs and attorney fees clause')
    
    doc.add_page_break()
    
    # Liability Protection
    add_heading_custom(doc, 'LIABILITY PROTECTION PROVISIONS', level=1)
    
    add_paragraph_custom(doc, 'Required Insurance (Contractor-Paid):', bold=True)
    
    add_paragraph_custom(doc, '1. Errors & Omissions (E&O) Insurance', bold=True)
    add_paragraph_custom(doc, '• Minimum: $1,000,000 per occurrence / $2,000,000 aggregate')
    add_paragraph_custom(doc, '• Coverage for: Misrepresentation, failure to disclose, negligent advice')
    add_paragraph_custom(doc, '• FDHC named as additional insured')
    add_paragraph_custom(doc, '• Certificate of insurance provided before first sale')
    
    add_paragraph_custom(doc, '2. General Liability Insurance', bold=True)
    add_paragraph_custom(doc, '• Minimum: $1,000,000 per occurrence')
    add_paragraph_custom(doc, '• Coverage for: Property damage, bodily injury during sales activities')
    add_paragraph_custom(doc, '• FDHC named as additional insured')
    
    add_paragraph_custom(doc, '3. Auto Liability (if using personal vehicle)', bold=True)
    add_paragraph_custom(doc, '• Minimum: $100,000/$300,000/$100,000')
    add_paragraph_custom(doc, '• Business use endorsement required')
    
    doc.add_paragraph()
    add_paragraph_custom(doc, 'Contractual Protections:', bold=True)
    
    add_paragraph_custom(doc, '1. Independent Contractor Status Acknowledgment', bold=True)
    add_paragraph_custom(doc, 'Contractor acknowledges and agrees that:')
    add_paragraph_custom(doc, '(a) They are an independent contractor, not an employee of FDHC', indent=True)
    add_paragraph_custom(doc, '(b) They are responsible for all federal, state, and local taxes', indent=True)
    add_paragraph_custom(doc, '(c) They will receive Form 1099-NEC, not Form W-2', indent=True)
    add_paragraph_custom(doc, '(d) They are not entitled to employee benefits', indent=True)
    add_paragraph_custom(doc, '(e) They set their own schedule and methods', indent=True)
    add_paragraph_custom(doc, '(f) They may perform services for other companies, including competitors', indent=True)
    
    doc.add_page_break()
    
    add_paragraph_custom(doc, '2. Hold Harmless Clause', bold=True)
    add_paragraph_custom(doc, 'Contractor shall indemnify, defend, and hold harmless FDHC, its officers, directors, employees, and agents from and against any and all claims, damages, losses, liabilities, costs, and expenses (including reasonable attorney fees) arising out of or relating to:')
    add_paragraph_custom(doc, '(a) Contractor\'s performance under this Agreement', indent=True)
    add_paragraph_custom(doc, '(b) Any negligent or wrongful act or omission of Contractor', indent=True)
    add_paragraph_custom(doc, '(c) Any breach of this Agreement by Contractor', indent=True)
    add_paragraph_custom(doc, '(d) Any violation of applicable law by Contractor', indent=True)
    
    add_paragraph_custom(doc, '3. Limitation of FDHC Liability', bold=True)
    add_paragraph_custom(doc, 'FDHC\'s liability to Contractor for any claim arising out of this Agreement shall not exceed the total commissions paid to Contractor in the 12 months immediately preceding the claim. FDHC shall not be liable for any indirect, incidental, special, consequential, or punitive damages.')
    
    add_paragraph_custom(doc, '4. No Authority to Bind', bold=True)
    add_paragraph_custom(doc, 'Contractor has no authority to bind FDHC to any obligation, make any representation or warranty on behalf of FDHC, or enter into any agreement on behalf of FDHC. All representations to customers shall clearly indicate Contractor\'s independent status.')
    
    add_paragraph_custom(doc, '5. Customer Disclosure Requirement', bold=True)
    add_paragraph_custom(doc, 'Contractor shall provide each customer with a written disclosure stating: "[Contractor Name] is an independent contractor and not an employee of Factory Direct Homes Center. All representations are made solely by the independent contractor and not by Factory Direct Homes Center."')
    
    doc.add_page_break()
    
    add_heading_custom(doc, 'Auburn/DeKalb County-Specific Considerations', level=2)
    add_paragraph_custom(doc, 'Local Business License:')
    add_paragraph_custom(doc, '• DeKalb County may require contractor to obtain business license')
    add_paragraph_custom(doc, '• Check with Auburn City Clerk: (260) 925-1234')
    add_paragraph_custom(doc, '• Cost: Typically $25-50 annually')
    
    add_paragraph_custom(doc, 'Sales Tax:')
    add_paragraph_custom(doc, '• Contractor is NOT responsible for sales tax collection')
    add_paragraph_custom(doc, '• FDHC remains responsible for all tax collection and remittance')
    add_paragraph_custom(doc, '• Clearly document in agreement')
    
    add_paragraph_custom(doc, 'Zoning/Permitting:')
    add_paragraph_custom(doc, '• Contractor should NOT assist with permitting (creates employee-like control)')
    add_paragraph_custom(doc, '• FDHC or third-party handles all permitting')
    add_paragraph_custom(doc, '• Contractor may provide information only')
    
    doc.add_page_break()
    
    # Documentation Requirements
    add_heading_custom(doc, 'DOCUMENTATION REQUIREMENTS', level=1)
    
    add_paragraph_custom(doc, 'Before First Sale:', bold=True)
    add_paragraph_custom(doc, '1. Signed Independent Contractor Agreement (attached template)')
    add_paragraph_custom(doc, '2. W-9 Form (IRS Form, kept on file)')
    add_paragraph_custom(doc, '3. Certificate of Insurance (E&O + GL, FDHC named as AI)')
    add_paragraph_custom(doc, '4. Business License (if required by DeKalb County)')
    add_paragraph_custom(doc, '5. Disclosure Acknowledgment (customer disclosure requirement)')
    
    add_paragraph_custom(doc, 'Ongoing:', bold=True)
    add_paragraph_custom(doc, '1. Monthly Commission Statements (detailed breakdown)')
    add_paragraph_custom(doc, '2. 1099-NEC Filed by January 31 (for prior year payments)')
    add_paragraph_custom(doc, '3. Lead Source Documentation (for self-generated lead bonuses)')
    add_paragraph_custom(doc, '4. Customer Disclosure Forms (retained in customer file)')
    
    doc.add_page_break()
    
    # Comparison Summary
    add_heading_custom(doc, 'COMPARISON SUMMARY', level=1)
    add_table_custom(doc,
        ['Feature', 'Plan 1: Straight Commission', 'Plan 2: Tiered Performance', 'Plan 3: Draw Model'],
        [
            ['IRS Risk', 'LOW', 'LOW', 'HIGH'],
            ['Contractor Income Stability', 'LOW', 'MEDIUM', 'HIGH'],
            ['FDHC Cash Flow', 'BEST', 'GOOD', 'FAIR'],
            ['Down Payment Options', 'All 3 options', 'Milestone/Bonus only', 'None recommended'],
            ['Best For', 'Experienced self-starters', 'Growth-focused contractors', 'New contractors (with legal review)'],
            ['Admin Complexity', 'LOW', 'MEDIUM', 'HIGH'],
            ['Commission %', '2-3%', '2.5-3.5%', '3-4%']
        ]
    )
    
    doc.add_page_break()
    
    # Recommendations
    add_heading_custom(doc, 'RECOMMENDATIONS', level=1)
    
    add_paragraph_custom(doc, 'For FDHC:', bold=True)
    add_paragraph_custom(doc, '1. Start with Plan 1 or 2 - Lower legal risk, easier administration')
    add_paragraph_custom(doc, '2. Require E&O Insurance - Non-negotiable for all contractors')
    add_paragraph_custom(doc, '3. Document Everything - Lead sources, customer disclosures, commission calculations')
    add_paragraph_custom(doc, '4. Annual Legal Review - Have employment attorney review classification annually')
    add_paragraph_custom(doc, '5. Avoid Plan 3 - Unless absolutely necessary and with extensive legal documentation')
    
    add_paragraph_custom(doc, 'For 1099 Compliance:', bold=True)
    add_paragraph_custom(doc, '1. Never set hours - Contractors set their own schedule')
    add_paragraph_custom(doc, '2. Never require exclusivity - They can work for competitors')
    add_paragraph_custom(doc, '3. Never provide benefits - No health insurance, retirement, PTO')
    add_paragraph_custom(doc, '4. Never reimburse expenses - They pay their own way')
    add_paragraph_custom(doc, '5. Never mandate methods - They choose how to sell')
    add_paragraph_custom(doc, '6. Always pay commission-only - No base, no guarantee, no salary')
    
    add_paragraph_custom(doc, 'Down Payment Best Practice:', bold=True)
    add_paragraph_custom(doc, 'Use Option C (Milestone Payment) from Plan 1:')
    add_paragraph_custom(doc, '• 25% at contract signing (with down payment)')
    add_paragraph_custom(doc, '• Clearly tied to customer commitment')
    add_paragraph_custom(doc, '• Documented as performance milestone')
    add_paragraph_custom(doc, '• Lowest IRS risk while providing contractor cash flow')
    
    doc.add_page_break()
    
    # Legal Disclaimer
    add_heading_custom(doc, 'LEGAL DISCLAIMER', level=1)
    
    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run('IMPORTANT: ')
    disclaimer_run.bold = True
    disclaimer_run.font.color.rgb = RGBColor(255, 0, 0)
    disclaimer_para.add_run('This document is for informational purposes only and does not constitute legal advice. Worker classification is fact-specific and determined by the totality of the working relationship, not just the compensation structure.')
    
    add_paragraph_custom(doc, 'Before implementing any 1099 sales program, FDHC should:', bold=True)
    add_paragraph_custom(doc, '1. Consult with an Indiana-licensed employment attorney')
    add_paragraph_custom(doc, '2. Have all agreements reviewed for compliance with current law')
    add_paragraph_custom(doc, '3. Consider requesting an IRS SS-8 Determination for the first contractor')
    add_paragraph_custom(doc, '4. Review classification annually as laws and enforcement priorities change')
    
    add_paragraph_custom(doc, 'Misclassification Penalties (Federal):', bold=True)
    add_paragraph_custom(doc, '• Failure to withhold income tax: 1.5% of wages + 100% of FICA')
    add_paragraph_custom(doc, '• Willful misclassification: Up to $1,000 per worker + criminal penalties')
    add_paragraph_custom(doc, '• State penalties vary (Indiana: back taxes + interest + potential fines)')
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Document Version: 1.0 | Last Updated: April 16, 2026 | Next Review: April 16, 2027')
    footer_run.font.name = 'Calibri'
    footer_run.font.size = Pt(9)
    footer_run.italic = True
    
    # Save document
    output_path = '/Users/ava/.openclaw/workspace/FDHC_1099_Sales_Pay_Plans.docx'
    doc.save(output_path)
    print(f'Document saved successfully to: {output_path}')

if __name__ == '__main__':
    main()